"""Requirements_TR.docx üretir - Phase 1 teslim icin Turkce versiyon.

Calistirma:
    python scripts/make_requirements_tr_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

from group_info import GROUP_TR

PROJECT_DIR = Path(__file__).resolve().parent.parent
OUTPUT = PROJECT_DIR / "Requirements_TR.docx"


def add_para(doc, text="", *, style=None, bold=False, italic=False,
             size=None, align=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        if bold:
            r.bold = True
        if italic:
            r.italic = True
        if size:
            r.font.size = Pt(size)
    return p


def add_runs(doc, parts, *, style=None, align=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    for item in parts:
        if isinstance(item, str):
            p.add_run(item)
            continue
        text, attrs = item
        r = p.add_run(text)
        if attrs.get("bold"):
            r.bold = True
        if attrs.get("italic"):
            r.italic = True
        if attrs.get("mono"):
            r.font.name = "Consolas"
            r.font.size = Pt(10)
    return p


def shade_cell(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_table(doc, headers, rows, *, header_fill="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], header_fill)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        for ci, cell_parts in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            if isinstance(cell_parts, str):
                cell_parts = [(cell_parts, {})]
            for text, attrs in cell_parts:
                rr = p.add_run(text)
                if attrs.get("bold"):
                    rr.bold = True
                if attrs.get("italic"):
                    rr.italic = True
                if attrs.get("mono"):
                    rr.font.name = "Consolas"
                    rr.font.size = Pt(10)
    return table


def set_base_font(doc, family="Calibri", size=11):
    style = doc.styles["Normal"]
    style.font.name = family
    style.font.size = Pt(size)


def build() -> None:
    doc = Document()
    set_base_font(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ---- Baslik ----
    add_para(doc, "COM2058 – Proje Faz 1",
             style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Formula 1 Veritabanı – Veri Gereksinimleri",
             style="Heading 1", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_runs(doc, [
        ("Grup: ", {"bold": True}),
        (GROUP_TR, {}),
    ])
    add_runs(doc, [
        ("Tarih: ", {"bold": True}),
        ("26.04.2026", {}),
    ])
    add_runs(doc, [
        ("Uygulama alanı: ", {"bold": True}),
        ("Formula 1 Dünya Şampiyonası istatistik sistemi.", {}),
    ])

    add_para(doc)

    note = doc.add_paragraph()
    note.paragraph_format.left_indent = Cm(0.6)
    note.paragraph_format.right_indent = Cm(0.6)
    r = note.add_run(
        "Bu döküman, derste işlenen Elmasri & Navathe – "
        "Fundamentals of Database Systems (6. baskı), Bölüm 7'deki COMPANY "
        "veritabanı örneğinin anlatım tarzında hazırlanmıştır. "
        "Formula 1 bilgi sisteminin modellendiği mini-dünyayı açıklar; "
        "sonraki fazlarda bu tanıma uygun olarak şema ve uygulama "
        "gerçekleştirilecektir."
    )
    r.italic = True

    # ---- 1. Mini-dunya tanimi ----
    add_para(doc, "1. Mini-dünya tanımı", style="Heading 1")

    add_para(doc,
        "Veritabanı, birden çok sezon boyunca Formula 1 Dünya Şampiyonası'nı; "
        "takımları (konstrüktörler), pilotları, pistleri ve yarışları ile her "
        "yarış sonucunu tutar. Ana veri gereksinimleri aşağıdaki gibidir:")

    # 1.1 seasons
    add_runs(doc, [
        ("(1) ", {"bold": True}),
        ("Şampiyona ", {}),
        ("sezonlar", {"bold": True}),
        (" halinde düzenlenir. Her sezon, ", {}),
        ("yıl", {"bold": True}),
        (" bilgisiyle tekil olarak belirlenir. Bir sezonun 1950 yılı veya "
         "sonrası olması gerekir (ilk Formula 1 şampiyonası). Her sezon çok "
         "sayıda yarış içerir ve tek bir konstrüktör sıralaması üretir.", {}),
    ])

    # 1.2 teams
    add_runs(doc, [
        ("(2) ", {"bold": True}),
        ("Bir ", {}),
        ("takım", {"bold": True}),
        (" (konstrüktör) şampiyonaya katılır. Her takım için ", {}),
        ("takım id'si", {"bold": True}),
        (" (tekil, vekil anahtar), tekil bir ", {}),
        ("kısa isim", {"bold": True}),
        (" (ör. Ferrari), uzun ", {}),
        ("tam isim", {"bold": True}),
        (" (ör. Scuderia Ferrari HP), ", {}),
        ("kuruluş yılı", {"bold": True}),
        (", karargâhın bulunduğu ", {}),
        ("ülke", {"bold": True}),
        (" ve ", {}),
        ("şehir", {"bold": True}),
        (", güncel ", {}),
        ("takım şefi", {"bold": True}),
        (" adı ve güncel ", {}),
        ("motor tedarikçisi", {"bold": True}),
        (" tutulur.", {}),
    ])

    # 1.3 drivers
    add_runs(doc, [
        ("(3) ", {"bold": True}),
        ("Bir ", {}),
        ("pilot", {"bold": True}),
        (" şampiyonada yarışır. Her pilot için ", {}),
        ("pilot id'si", {"bold": True}),
        (" (tekil), yarıştığı ", {}),
        ("araç numarası", {"bold": True}),
        (" (atandığında tekil), üç harfli ", {}),
        ("pilot kodu", {"bold": True}),
        (" (tekil, ör. HAM, VER), ", {}),
        ("ad", {"bold": True}),
        (", ", {}),
        ("soyad", {"bold": True}),
        (", ", {}),
        ("uyruk", {"bold": True}),
        (", ", {}),
        ("doğum tarihi", {"bold": True}),
        (" ve pilotun ", {}),
        ("kariyer istatistikleri", {"bold": True}),
        (" saklanır: pole pozisyonu sayısı, podyum sayısı ve kazandığı "
         "şampiyonluk sayısı. Bu üç sayaç ", {}),
        ("türetilmiş nitelik", {"bold": True}),
        (" olarak değerlendirilir — yarış sonuçları ve konstrüktör "
         "sıralamalarından tekrar hesaplanabilir, ancak pilot "
         "kartlarını gösteren web sayfalarını hızlandırmak için "
         "tabloda saklanır.", {}),
    ])

    # 1.4 circuits
    add_runs(doc, [
        ("(4) ", {"bold": True}),
        ("Bir ", {}),
        ("pist", {"bold": True}),
        (" Grand Prix etkinliklerine ev sahipliği yapar. Her pist için ", {}),
        ("pist id'si", {"bold": True}),
        (" (tekil), tekil ", {}),
        ("pist adı", {"bold": True}),
        (" (ör. Circuit de Monaco), bulunduğu ", {}),
        ("ülke", {"bold": True}),
        (" ve ", {}),
        ("şehir", {"bold": True}),
        (", normal yarış mesafesinin ", {}),
        ("tur sayısı", {"bold": True}),
        (", metre cinsinden ", {}),
        ("pist uzunluğu", {"bold": True}),
        (", ", {}),
        ("ilk Grand Prix yılı", {"bold": True}),
        (" ve o pistte kaydedilen mevcut ", {}),
        ("tüm zamanların en hızlı tur süresi", {"bold": True}),
        (" tutulur.", {}),
    ])

    # 1.5 races
    add_runs(doc, [
        ("(5) ", {"bold": True}),
        ("Bir ", {}),
        ("yarış", {"bold": True}),
        (", bir sezon içinde belirli bir pistte gerçekleşen Grand Prix "
         "etkinliğidir. Bir yarış, ", {}),
        ("sezon yılı", {"bold": True}),
        (" ile o sezonun içindeki ", {}),
        ("ayak numarası", {"bold": True}),
        (" birlikte tekil olarak belirlenir; dolayısıyla yarış kendi "
         "sezonu olmadan var olamaz ve SEASON'a ", {}),
        ("HAS_RACE", {"italic": True}),
        (" ilişkisiyle bağlı bir ", {}),
        ("zayıf varlık (weak entity)", {"bold": True}),
        (" olarak modellenir. Her yarış için ayrıca ", {}),
        ("Grand Prix adı", {"bold": True}),
        (" (bir sezon içinde tekil, ör. Monaco Grand Prix), yarış "
         "haftasonunun ", {}),
        ("başlangıç", {"bold": True}),
        (" ve ", {}),
        ("bitiş tarihi", {"bold": True}),
        (" tutulur. Her yarış tam olarak ", {}),
        ("bir", {"bold": True}),
        (" piste ev sahipliği yapar; bir pist ise yıllar içinde ", {}),
        ("birçok", {"bold": True}),
        (" yarışa (", {}),
        ("HOSTS", {"italic": True}),
        (" ilişkisi) ev sahipliği yapabilir.", {}),
    ])

    # 1.6 contracts
    add_runs(doc, [
        ("(6) ", {"bold": True}),
        ("Her sezonda pilotlar takımlarla ", {}),
        ("sözleşme yapar", {"bold": True}),
        (". Her atama bir sezon, bir takım ve bir pilot olmak üzere üç "
         "varlıkla birlikte (üçlü ilişki ", {}),
        ("CONTRACTS", {"italic": True}),
        (") tanımlanır. Her sözleşme için pilotun ", {}),
        ("rolü", {"bold": True}),
        (" — MAIN (yarışan pilot) veya RESERVE (yedek) — ve yalnızca MAIN "
         "pilotlar için bir ", {}),
        ("koltuk numarası", {"bold": True}),
        (" (1 veya 2) tutulur. İşleme kuralı şöyledir:", {}),
    ])
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("eğer ")
    r2 = p.add_run("role = MAIN"); r2.font.name = "Consolas"; r2.font.size = Pt(10)
    p.add_run("  →  ")
    r3 = p.add_run("seat_no ∈ {1, 2}"); r3.font.name = "Consolas"; r3.font.size = Pt(10)
    p.add_run(" (NOT NULL)")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("eğer ")
    r4 = p.add_run("role = RESERVE"); r4.font.name = "Consolas"; r4.font.size = Pt(10)
    p.add_run("  →  ")
    r5 = p.add_run("seat_no IS NULL"); r5.font.name = "Consolas"; r5.font.size = Pt(10)

    # 1.7 standings
    add_runs(doc, [
        ("(7) ", {"bold": True}),
        ("Her sezonun sonunda, o sezona katılan her takım için bir ", {}),
        ("konstrüktör sıralaması", {"bold": True}),
        (" üretilir. Bu sıralama, SEASON ve TEAM arasında ", {}),
        ("final pozisyonu", {"bold": True}),
        (" ve sezonda kazanılan ", {}),
        ("toplam puan", {"bold": True}),
        (" niteliklerini taşıyan ikili bir ilişkidir (", {}),
        ("STANDINGS", {"italic": True}),
        (").", {}),
    ])

    # 1.8 results
    add_runs(doc, [
        ("(8) ", {"bold": True}),
        ("Her yarış için, yarışa katılan her pilotun ", {}),
        ("yarış sonucu", {"bold": True}),
        (" kaydedilir. Bu, RACE ve DRIVER arasında ikili bir ilişkidir (", {}),
        ("RESULTS", {"italic": True}),
        (") ve nitelikleri pilotun yarış başlangıcındaki ", {}),
        ("grid pozisyonu", {"bold": True}),
        (", ", {}),
        ("bitiriş pozisyonu", {"bold": True}),
        (" ve o yarışta kazanılan şampiyona ", {}),
        ("puanı", {"bold": True}),
        ("dır.", {}),
    ])

    # ---- 2. Fonksiyonel gereksinimler ----
    add_para(doc, "2. Fonksiyonel gereksinimler (desteklenmesi gereken sorgular)",
             style="Heading 1")
    add_para(doc, "İlişkisel şema, aşağıdaki örnek sorguları verimli bir "
                  "biçimde destekleyebilmelidir:")
    queries = [
        "Verilen bir sezonun tüm yarışlarını pist adı ve Grand Prix adıyla "
        "birlikte kronolojik sırayla listele.",
        "Verilen bir sezon için konstrüktör sıralamasını "
        "(takım → pozisyon, puan) pozisyona göre sıralı olarak getir.",
        "Verilen bir pilot için katıldığı tüm yarışları; başlangıç grid "
        "pozisyonu, bitiriş pozisyonu ve kazanılan puan ile listele.",
        "Verilen bir takımın verilen bir sezonda sözleşmeli olduğu tüm "
        "pilotları rol ve koltuk numarası ile listele.",
        "Verilen bir pistin şimdiye kadar ev sahipliği yaptığı tüm Grand Prix "
        "etkinliklerini listele.",
        "Verilen bir pilotun kariyer toplamlarını (şampiyonluk, podyum, pole) "
        "getir — bu değerler türetilmiş nitelikler olarak tutulur.",
    ]
    for q in queries:
        doc.add_paragraph(q, style="List Number")

    # ---- 3. Ozet ----
    add_para(doc, "3. Varlık tipleri, ilişki tipleri ve kısıt özeti",
             style="Heading 1")

    # 3.1
    add_para(doc, "3.1  Varlık tipleri (entity types)", style="Heading 2")
    add_table(doc,
        headers=["Varlık", "Tür", "Anahtar nitelik(ler)"],
        rows=[
            ["SEASON", "güçlü", [("season_year", {"mono": True})]],
            ["TEAM", "güçlü",
                [("team_id", {"mono": True}), (" (PK), ", {}),
                 ("team_name", {"mono": True}), (" (UK)", {})]],
            ["DRIVER", "güçlü",
                [("driver_id", {"mono": True}), (" (PK), ", {}),
                 ("car_no", {"mono": True}), (" (UK), ", {}),
                 ("driver_code", {"mono": True}), (" (UK)", {})]],
            ["CIRCUIT", "güçlü",
                [("circuit_id", {"mono": True}), (" (PK), ", {}),
                 ("circuit_name", {"mono": True}), (" (UK)", {})]],
            ["RACE",
                [("zayıf (weak)", {"bold": True}),
                 (" — SEASON tarafından tanımlanır", {})],
                [("kısmi anahtar (partial key): ", {}),
                 ("round_number", {"mono": True}),
                 (" (", {}), ("season_year", {"mono": True}),
                 (" ile birlikte)", {})]],
        ])

    # 3.2
    add_para(doc, "3.2  İlişki tipleri (relationship types)", style="Heading 2")
    add_table(doc,
        headers=["İlişki", "Derece", "Katılan varlıklar",
                 "Kardinalite", "İlişki nitelikleri"],
        rows=[
            ["HAS_RACE",
                [("ikili, ", {}), ("tanımlayıcı (identifying)", {"bold": True})],
                "SEASON — RACE",
                "1 : N (RACE tam katılım)", "—"],
            ["HOSTS", "ikili", "CIRCUIT — RACE",
                "1 : N (RACE tam katılım)", "—"],
            ["CONTRACTS",
                [("üçlü (ternary)", {"bold": True})],
                "SEASON — TEAM — DRIVER", "M : N : P",
                [("role", {"mono": True}), (", ", {}),
                 ("seat_no", {"mono": True})]],
            ["STANDINGS", "ikili", "SEASON — TEAM", "M : N",
                [("season_position", {"mono": True}), (", ", {}),
                 ("season_points", {"mono": True})]],
            ["RESULTS", "ikili", "RACE — DRIVER", "M : N",
                [("grid_position", {"mono": True}), (", ", {}),
                 ("finish_position", {"mono": True}), (", ", {}),
                 ("points", {"mono": True})]],
        ])

    # 3.3
    add_para(doc, "3.3  Anahtar ve bütünlük kuralları", style="Heading 2")
    rules = [
        [("season_year", {"mono": True}), (" ≥ 1950", {})],
        [("team_name", {"mono": True}), (" tekil; ", {}),
         ("driver_code", {"mono": True}), (" tekil; ", {}),
         ("car_no", {"mono": True}), (" tekil (atandığında); ", {}),
         ("circuit_name", {"mono": True}), (" tekil.", {})],
        [("Yarış: ", {"bold": True}),
         ("(season_year, round_number)", {"mono": True}),
         (" birleşik birincil anahtardır; ", {}),
         ("round_number > 0", {"mono": True}), ("; ", {}),
         ("race_start_date ≤ race_end_date", {"mono": True}), ("; ", {}),
         ("Grand Prix adı bir sezon içinde tekildir.", {})],
        [("Sonuçlar: ", {"bold": True}),
         ("grid_position > 0", {"mono": True}), (", ", {}),
         ("finish_position > 0", {"mono": True}), (", ", {}),
         ("points ≥ 0", {"mono": True}), (".", {})],
        [("Sıralama: ", {"bold": True}),
         ("season_points ≥ 0", {"mono": True}), (".", {})],
        [("Pilot kariyer sayaçları: ", {"bold": True}),
         ("num_poles ≥ 0", {"mono": True}), (", ", {}),
         ("num_podiums ≥ 0", {"mono": True}), (", ", {}),
         ("num_championships ≥ 0", {"mono": True}), (".", {})],
    ]
    for parts in rules:
        p = doc.add_paragraph(style="List Bullet")
        for text, attrs in parts:
            rr = p.add_run(text)
            if attrs.get("bold"):
                rr.bold = True
            if attrs.get("mono"):
                rr.font.name = "Consolas"
                rr.font.size = Pt(10)

    # 3.4
    add_para(doc, "3.4  Türetilmiş nitelikler (derived attributes)",
             style="Heading 2")
    add_runs(doc, [
        ("DRIVER.num_poles", {"mono": True}), (", ", {}),
        ("num_podiums", {"mono": True}), (", ", {}),
        ("num_championships", {"mono": True}),
        (" nitelikleri ", {}),
        ("RESULTS", {"mono": True}), (" ve ", {}),
        ("STANDINGS", {"mono": True}),
        (" ilişkilerinden türetilir. Performans için tabloda saklanır ve "
         "uygulama katmanında güncel tutulur.", {}),
    ])

    # ---- 4. Sonraki fazlar ----
    add_para(doc, "4. Sonraki fazlar için teslimler", style="Heading 1")
    deliverables = [
        [("Faz 2 – ER Diyagramı (Chen gösterimi): ", {"bold": True}),
         ("ERD.drawio", {"mono": True}), (" / ", {}),
         ("ERD.pdf", {"mono": True}),
         (" (bu dokümanla birlikte teslim edilir).", {})],
        [("Faz 3 – Uygulama: ", {"bold": True}),
         ("MySQL 8 şeması (", {}),
         ("Formula_1.sql", {"mono": True}),
         (") + FastAPI backend (yalnızca ham SQL, ORM yok) + FastAPI "
          "tarafından sunulan SPA frontend.", {})],
        [("Faz 4 – Rapor (10–15 sayfa): ", {"bold": True}),
         ("veri gereksinimleri, ER modeli, ilişkisel modele dönüşüm, "
          "normalizasyon (3NF / BCNF'ye kadar), şema, örnek sorgular ve "
          "ekran görüntülerini kapsar.", {})],
    ]
    for parts in deliverables:
        p = doc.add_paragraph(style="List Bullet")
        for text, attrs in parts:
            rr = p.add_run(text)
            if attrs.get("bold"):
                rr.bold = True
            if attrs.get("mono"):
                rr.font.name = "Consolas"
                rr.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
