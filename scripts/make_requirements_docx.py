"""Generate Requirements.docx from Requirements.md for Phase 1 submission.

Runs with python-docx. Usage:
    python scripts/make_requirements_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm


PROJECT_DIR = Path(__file__).resolve().parent.parent
OUTPUT = PROJECT_DIR / "Requirements.docx"


# -----------------------------
# Small helpers
# -----------------------------
def add_para(doc, text="", *, style=None, bold=False, italic=False,
             size=None, align=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        if bold:
            r.bold = True
        if italic:
            r.italic = True
        if size:
            r.font.size = Pt(size)
    return p


def add_runs(doc, parts, *, style=None, align=None):
    """Add a paragraph with a list of (text, {bold, italic, mono}) tuples."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    for item in parts:
        if isinstance(item, str):
            p.add_run(item)
            continue
        text, attrs = item
        r = p.add_run(text)
        if attrs.get("bold"):
            r.bold = True
        if attrs.get("italic"):
            r.italic = True
        if attrs.get("mono"):
            r.font.name = "Consolas"
            r.font.size = Pt(10)
    return p


def shade_cell(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_table(doc, headers, rows, *, header_fill="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], header_fill)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        for ci, cell_parts in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            if isinstance(cell_parts, str):
                cell_parts = [(cell_parts, {})]
            for text, attrs in cell_parts:
                rr = p.add_run(text)
                if attrs.get("bold"):
                    rr.bold = True
                if attrs.get("italic"):
                    rr.italic = True
                if attrs.get("mono"):
                    rr.font.name = "Consolas"
                    rr.font.size = Pt(10)
    return table


def set_base_font(doc, family="Calibri", size=11):
    style = doc.styles["Normal"]
    style.font.name = family
    style.font.size = Pt(size)


# -----------------------------
# Build document
# -----------------------------
def build() -> None:
    doc = Document()
    set_base_font(doc)

    # Narrow margins for a cleaner 1-page look
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ---- Title block ----
    add_para(doc, "COM2058 – Project Phase 1",
             style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Formula 1 Database – Data Requirements",
             style="Heading 1", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_runs(doc, [
        ("Group: ", {"bold": True}),
        ("TODO: names & student IDs", {"italic": True}),
    ])
    add_runs(doc, [
        ("Date: ", {"bold": True}),
        ("26.04.2026", {}),
    ])
    add_runs(doc, [
        ("Application domain: ", {"bold": True}),
        ("Formula 1 World Championship statistics system.", {}),
    ])

    add_para(doc)

    note = doc.add_paragraph()
    note.paragraph_format.left_indent = Cm(0.6)
    note.paragraph_format.right_indent = Cm(0.6)
    r = note.add_run(
        "This document is written in the same narrative style as the COMPANY "
        "database example in Elmasri & Navathe – Fundamentals of Database "
        "Systems (6e), Chapter 7, which was presented in class. It describes "
        "the mini-world of the Formula 1 information system that will be "
        "implemented in the following phases."
    )
    r.italic = True

    # ---- 1. Mini-world description ----
    add_para(doc, "1. Mini-world description", style="Heading 1")

    add_para(doc,
        "The database keeps track of the Formula 1 World Championship across "
        "multiple seasons, teams (constructors), drivers, circuits and races, "
        "as well as the results of every race weekend. The main data "
        "requirements are the following:")

    # 1.1 seasons
    add_runs(doc, [
        ("(1) ", {"bold": True}),
        ("The championship is organised into ", {}),
        ("seasons", {"bold": True}),
        (". Each season is uniquely identified by its ", {}),
        ("year", {"bold": True}),
        (". A season must take place in 1950 or later (the first Formula 1 "
         "championship). A season consists of many races and produces one "
         "constructors' standings table.", {}),
    ])

    # 1.2 teams
    add_runs(doc, [
        ("(2) ", {"bold": True}),
        ("A ", {}),
        ("team", {"bold": True}),
        (" (constructor) participates in the championship. For every team we "
         "store a ", {}),
        ("team id", {"bold": True}),
        (" (unique, surrogate), a ", {}),
        ("short name", {"bold": True}),
        (" (unique, e.g. Ferrari), a long ", {}),
        ("full name", {"bold": True}),
        (" (e.g. Scuderia Ferrari HP), the ", {}),
        ("founding year", {"bold": True}),
        (", the ", {}),
        ("base country", {"bold": True}),
        (" and ", {}),
        ("base city", {"bold": True}),
        (" of its headquarters, the name of the current ", {}),
        ("team chief", {"bold": True}),
        (", and the current ", {}),
        ("engine supplier", {"bold": True}),
        (".", {}),
    ])

    # 1.3 drivers
    add_runs(doc, [
        ("(3) ", {"bold": True}),
        ("A ", {}),
        ("driver", {"bold": True}),
        (" competes in the championship. For every driver we store a ", {}),
        ("driver id", {"bold": True}),
        (" (unique), the ", {}),
        ("car number", {"bold": True}),
        (" the driver races with (unique when assigned), a three-letter ", {}),
        ("driver code", {"bold": True}),
        (" (unique, e.g. HAM, VER), ", {}),
        ("first name", {"bold": True}),
        (", ", {}),
        ("last name", {"bold": True}),
        (", ", {}),
        ("nationality", {"bold": True}),
        (", ", {}),
        ("date of birth", {"bold": True}),
        (", and the driver's ", {}),
        ("career statistics", {"bold": True}),
        (": number of pole positions, number of podium finishes, and number "
         "of championships won. These three career counters are treated as ", {}),
        ("derived attributes", {"bold": True}),
        (" — they can be recomputed from race results and constructors' "
         "standings, but are stored to speed up the web pages that show "
         "driver cards.", {}),
    ])

    # 1.4 circuits
    add_runs(doc, [
        ("(4) ", {"bold": True}),
        ("A ", {}),
        ("circuit", {"bold": True}),
        (" (track) hosts Grand Prix events. For every circuit we store a ", {}),
        ("circuit id", {"bold": True}),
        (" (unique), a ", {}),
        ("circuit name", {"bold": True}),
        (" (unique, e.g. Circuit de Monaco), the ", {}),
        ("country", {"bold": True}),
        (" and ", {}),
        ("city", {"bold": True}),
        (" it is located in, the ", {}),
        ("number of laps", {"bold": True}),
        (" of a normal race distance, the ", {}),
        ("track length", {"bold": True}),
        (" in metres, the ", {}),
        ("year of its first Grand Prix", {"bold": True}),
        (", and the current ", {}),
        ("all-time fastest lap time", {"bold": True}),
        (" recorded on the track.", {}),
    ])

    # 1.5 races
    add_runs(doc, [
        ("(5) ", {"bold": True}),
        ("A ", {}),
        ("race", {"bold": True}),
        (" is a Grand Prix event that takes place inside a season at a given "
         "circuit. A race is identified by the combination of its ", {}),
        ("season year", {"bold": True}),
        (" and its ", {}),
        ("round number", {"bold": True}),
        (" inside that season — therefore a race cannot exist without its "
         "season and is modelled as a ", {}),
        ("weak entity", {"bold": True}),
        (" identified by SEASON through the ", {}),
        ("HAS_RACE", {"italic": True}),
        (" relationship. For every race we also store the ", {}),
        ("Grand Prix name", {"bold": True}),
        (" (unique within a season, e.g. Monaco Grand Prix), the ", {}),
        ("start date", {"bold": True}),
        (" and ", {}),
        ("end date", {"bold": True}),
        (" of the race weekend. A race is hosted by ", {}),
        ("exactly one", {"bold": True}),
        (" circuit, and a circuit may host ", {}),
        ("many", {"bold": True}),
        (" races over the years (", {}),
        ("HOSTS", {"italic": True}),
        (" relationship).", {}),
    ])

    # 1.6 contracts
    add_runs(doc, [
        ("(6) ", {"bold": True}),
        ("In every season, drivers are ", {}),
        ("contracted", {"bold": True}),
        (" by teams. Each assignment is described by three entities together: "
         "a season, a team and a driver (a ternary relationship ", {}),
        ("CONTRACTS", {"italic": True}),
        ("). For every such contract we store the ", {}),
        ("role", {"bold": True}),
        (" of the driver — either MAIN (racing driver) or RESERVE — and, for "
         "main drivers only, a ", {}),
        ("seat number", {"bold": True}),
        (" (1 or 2). The business rule is therefore:", {}),
    ])
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("if "); r2 = p.add_run("role = MAIN"); r2.font.name = "Consolas"; r2.font.size = Pt(10)
    p.add_run("  →  ")
    r3 = p.add_run("seat_no ∈ {1, 2}"); r3.font.name = "Consolas"; r3.font.size = Pt(10)
    p.add_run(" (NOT NULL)")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("if "); r4 = p.add_run("role = RESERVE"); r4.font.name = "Consolas"; r4.font.size = Pt(10)
    p.add_run("  →  ")
    r5 = p.add_run("seat_no IS NULL"); r5.font.name = "Consolas"; r5.font.size = Pt(10)

    # 1.7 standings
    add_runs(doc, [
        ("(7) ", {"bold": True}),
        ("At the end of every season a ", {}),
        ("constructors' standing", {"bold": True}),
        (" is produced for every team that participated in that season. The "
         "standing is a binary relationship between SEASON and TEAM (", {}),
        ("STANDINGS", {"italic": True}),
        (") with attributes ", {}),
        ("final position", {"bold": True}),
        (" and ", {}),
        ("total points", {"bold": True}),
        (" earned in that season.", {}),
    ])

    # 1.8 results
    add_runs(doc, [
        ("(8) ", {"bold": True}),
        ("For every race, the ", {}),
        ("race result", {"bold": True}),
        (" of every driver that took part is recorded. This is a binary "
         "relationship between RACE and DRIVER (", {}),
        ("RESULTS", {"italic": True}),
        (") whose attributes are the driver's ", {}),
        ("grid position", {"bold": True}),
        (" at the start of the race, the ", {}),
        ("finishing position", {"bold": True}),
        (", and the number of championship ", {}),
        ("points", {"bold": True}),
        (" awarded for that race.", {}),
    ])

    # ---- 2. Functional requirements ----
    add_para(doc, "2. Functional requirements (queries that must be supported)",
             style="Heading 1")
    add_para(doc, "The relational schema must make the following example "
                  "queries efficient:")
    queries = [
        "List every race of a given season, in chronological order, together "
        "with the circuit name and the Grand Prix name.",
        "For a given season, return the constructors' standings (team → "
        "position, points), ordered by position.",
        "For a given driver, list every race they participated in, their "
        "starting grid position, finishing position and points earned.",
        "For a given team in a given season, list all the drivers contracted "
        "to that team with their role and seat number.",
        "For a given circuit, list all Grand Prix events it has hosted across "
        "all seasons.",
        "For a given driver, return the career totals (championships, podiums, "
        "poles) — which are maintained as derived attributes.",
    ]
    for q in queries:
        doc.add_paragraph(q, style="List Number")

    # ---- 3. Summary ----
    add_para(doc, "3. Summary of entity types, relationships and constraints",
             style="Heading 1")

    # 3.1
    add_para(doc, "3.1  Entity types", style="Heading 2")
    add_table(doc,
        headers=["Entity", "Kind", "Key attribute(s)"],
        rows=[
            ["SEASON", "strong", [("season_year", {"mono": True})]],
            ["TEAM", "strong",
                [("team_id", {"mono": True}), (" (PK), ", {}),
                 ("team_name", {"mono": True}), (" (UK)", {})]],
            ["DRIVER", "strong",
                [("driver_id", {"mono": True}), (" (PK), ", {}),
                 ("car_no", {"mono": True}), (" (UK), ", {}),
                 ("driver_code", {"mono": True}), (" (UK)", {})]],
            ["CIRCUIT", "strong",
                [("circuit_id", {"mono": True}), (" (PK), ", {}),
                 ("circuit_name", {"mono": True}), (" (UK)", {})]],
            ["RACE",
                [("weak", {"bold": True}), (" (identified by SEASON)", {})],
                [("partial key: ", {}), ("round_number", {"mono": True}),
                 (" (combined with ", {}), ("season_year", {"mono": True}),
                 (")", {})]],
        ])

    # 3.2
    add_para(doc, "3.2  Relationship types", style="Heading 2")
    add_table(doc,
        headers=["Relationship", "Degree", "Participating entities",
                 "Cardinality", "Relationship attributes"],
        rows=[
            ["HAS_RACE",
                [("binary, ", {}), ("identifying", {"bold": True})],
                "SEASON — RACE",
                "1 : N (RACE total participation)", "—"],
            ["HOSTS", "binary", "CIRCUIT — RACE",
                "1 : N (RACE total)", "—"],
            ["CONTRACTS",
                [("ternary", {"bold": True})],
                "SEASON — TEAM — DRIVER", "M : N : P",
                [("role", {"mono": True}), (", ", {}),
                 ("seat_no", {"mono": True})]],
            ["STANDINGS", "binary", "SEASON — TEAM", "M : N",
                [("season_position", {"mono": True}), (", ", {}),
                 ("season_points", {"mono": True})]],
            ["RESULTS", "binary", "RACE — DRIVER", "M : N",
                [("grid_position", {"mono": True}), (", ", {}),
                 ("finish_position", {"mono": True}), (", ", {}),
                 ("points", {"mono": True})]],
        ])

    # 3.3
    add_para(doc, "3.3  Key integrity rules", style="Heading 2")
    rules = [
        [("season_year", {"mono": True}), (" ≥ 1950", {})],
        [("team_name", {"mono": True}), (" unique; ", {}),
         ("driver_code", {"mono": True}), (" unique; ", {}),
         ("car_no", {"mono": True}), (" unique (when set); ", {}),
         ("circuit_name", {"mono": True}), (" unique.", {})],
        [("Race: ", {"bold": True}),
         ("(season_year, round_number)", {"mono": True}),
         (" is the composite primary key; ", {}),
         ("round_number > 0", {"mono": True}), ("; ", {}),
         ("race_start_date ≤ race_end_date", {"mono": True}), ("; ", {}),
         ("Grand Prix name unique within a season.", {})],
        [("Results: ", {"bold": True}),
         ("grid_position > 0", {"mono": True}), (", ", {}),
         ("finish_position > 0", {"mono": True}), (", ", {}),
         ("points ≥ 0", {"mono": True}), (".", {})],
        [("Standings: ", {"bold": True}),
         ("season_points ≥ 0", {"mono": True}), (".", {})],
        [("Driver career counters: ", {"bold": True}),
         ("num_poles ≥ 0", {"mono": True}), (", ", {}),
         ("num_podiums ≥ 0", {"mono": True}), (", ", {}),
         ("num_championships ≥ 0", {"mono": True}), (".", {})],
    ]
    for parts in rules:
        p = doc.add_paragraph(style="List Bullet")
        for text, attrs in parts:
            rr = p.add_run(text)
            if attrs.get("bold"):
                rr.bold = True
            if attrs.get("mono"):
                rr.font.name = "Consolas"
                rr.font.size = Pt(10)

    # 3.4
    add_para(doc, "3.4  Derived / computed attributes", style="Heading 2")
    add_runs(doc, [
        ("DRIVER.num_poles", {"mono": True}), (", ", {}),
        ("num_podiums", {"mono": True}), (", ", {}),
        ("num_championships", {"mono": True}),
        (" are derived from ", {}),
        ("RESULTS", {"mono": True}), (" and ", {}),
        ("STANDINGS", {"mono": True}),
        (". They are still stored in the table and kept in sync by the "
         "application layer.", {}),
    ])

    # ---- 4. Deliverables ----
    add_para(doc, "4. Deliverables for the following phases", style="Heading 1")
    deliverables = [
        [("Phase 2 – ER Diagram (Chen notation): ", {"bold": True}),
         ("ERD.drawio", {"mono": True}), (" / ", {}),
         ("ERD.pdf", {"mono": True}),
         (" (committed alongside this document).", {})],
        [("Phase 3 – Implementation: ", {"bold": True}),
         ("MySQL 8 schema (", {}),
         ("Formula_1.sql", {"mono": True}),
         (") + FastAPI backend (raw SQL only, no ORM) + SPA frontend served "
          "by FastAPI.", {})],
        [("Phase 4 – Report (10–15 pages): ", {"bold": True}),
         ("covers requirements, ER model, relational mapping, normalization "
          "(up to 3NF / BCNF), schema, selected queries and screenshots.", {})],
    ]
    for parts in deliverables:
        p = doc.add_paragraph(style="List Bullet")
        for text, attrs in parts:
            rr = p.add_run(text)
            if attrs.get("bold"):
                rr.bold = True
            if attrs.get("mono"):
                rr.font.name = "Consolas"
                rr.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
