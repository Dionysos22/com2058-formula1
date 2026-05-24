"""Generate REPORT.docx from REPORT.md (Phase 4 submission)."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from group_info import GROUP_EN

PROJECT_DIR = Path(__file__).resolve().parent.parent
REPORT_MD = PROJECT_DIR / "REPORT.md"
OUTPUT = PROJECT_DIR / "REPORT.docx"


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=min(level, 3))


def add_body(doc: Document, text: str) -> None:
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)


def add_table_from_md(doc: Document, lines: list[str]) -> None:
    rows = [line.strip().strip("|").split("|") for line in lines if line.strip().startswith("|")]
    if len(rows) < 2:
        return
    headers = [c.strip() for c in rows[0]]
    body = rows[2:] if len(rows) > 2 and set(rows[1]) <= {"-", ":", " "} else rows[1:]
    table = doc.add_table(rows=1 + len(body), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(body):
        for ci, cell in enumerate(row):
            table.rows[ri + 1].cells[ci].text = cell.strip()


def build() -> None:
    text = REPORT_MD.read_text(encoding="utf-8")
    text = text.replace("_TODO: names & student IDs_", GROUP_EN)
    text = text.replace("**Group:** _TODO: names & student IDs_", f"**Group:** {GROUP_EN}")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("COM2058 – Formula 1 Database Project\nPhase 4 Report")
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph(f"Group: {GROUP_EN}")
    doc.add_paragraph("Course: COM2058 – Database Systems")
    doc.add_paragraph("Date: 24.05.2026")
    doc.add_paragraph("")

    table_buf: list[str] = []
    in_code = False
    code_buf: list[str] = []

    for raw in text.splitlines():
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_body(doc, "\n".join(code_buf))
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        if line.startswith("|"):
            table_buf.append(line)
            continue
        if table_buf:
            add_table_from_md(doc, table_buf)
            table_buf = []

        if line.startswith("# "):
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 1)
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2)
            continue
        if line.strip() == "---":
            continue
        if line.startswith("**Group:**"):
            continue
        if line.startswith("**Course:**") or line.startswith("**Date:**") or line.startswith("**Application:**"):
            continue

        cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", line)
        cleaned = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", cleaned)
        cleaned = cleaned.replace("_", "")
        if cleaned.startswith("- "):
            doc.add_paragraph(cleaned[2:], style="List Bullet")
        elif cleaned.strip():
            add_body(doc, cleaned)

    if table_buf:
        add_table_from_md(doc, table_buf)

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
